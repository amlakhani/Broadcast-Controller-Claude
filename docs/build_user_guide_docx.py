import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Error: The 'python-docx' library is not installed.")
    print("Please install it by running: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "Broadcast_Controller_User_Guide.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x66, 0x77)
BLACK = RGBColor(0x11, 0x18, 0x27)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_paragraph_spacing(
        p,
        before=18 if level == 1 else 14,
        after=10 if level == 1 else 7
    )
    return p


def add_figure(doc, img_file, caption):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_file), width=Inches(6.25))
        set_paragraph_spacing(p, before=4, after=2)

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.runs[0]
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        set_paragraph_spacing(cap, after=10)
    except Exception as e:
        print(f"Warning: Could not insert image {img_file.name}: {e}")


def add_runs_to_paragraph(paragraph, text):
    # Split text by ** to handle markdown bold markers
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            if is_bold:
                run.bold = True
        is_bold = not is_bold


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK

    body = doc.styles["Body Text"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = BLACK
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def build():
    md_path = DOCS / "Broadcast_Controller_User_Guide.md"
    if not md_path.exists():
        print(f"Error: Markdown source file not found at {md_path}")
        sys.exit(1)

    print(f"Reading user guide from {md_path}...")
    doc = Document()
    style_document(doc)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    first_title = True
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Handle empty lines
        if not stripped:
            i += 1
            continue

        # 1. Horizontal dividers (---)
        if stripped in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("❖   ❖   ❖")
            run.font.color.rgb = MUTED
            set_paragraph_spacing(p, before=12, after=12)
            i += 1
            continue

        # 2. Images: ![Caption](Path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            caption = img_match.group(1)
            img_path_str = img_match.group(2)
            # Find the path relative to the DOCS directory
            img_file = DOCS / img_path_str
            if img_file.exists():
                add_figure(doc, img_file, caption)
            else:
                # Gracefully insert a placeholder callout block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                set_paragraph_spacing(p, before=6, after=6)
                run = p.add_run(f"[Visual Placeholder: {caption} ({img_path_str})]")
                run.italic = True
                run.font.color.rgb = MUTED
                run.font.size = Pt(10)
            i += 1
            continue

        # 3. Headings
        if stripped.startswith("# "):
            title_text = stripped[2:]
            if first_title:
                # Main Title styling
                title = doc.add_paragraph()
                title_run = title.add_run(title_text)
                title_run.font.name = "Calibri"
                title_run.font.size = Pt(24)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
                set_paragraph_spacing(title, before=10, after=4)
                first_title = False

                # Auto-detect subtitle if next non-empty line is not a heading
                next_idx = i + 1
                next_non_empty = None
                while next_idx < len(lines):
                    s = lines[next_idx].strip()
                    if s:
                        next_non_empty = (next_idx, s)
                        break
                    next_idx += 1

                if next_non_empty and not next_non_empty[1].startswith("#"):
                    subtitle = doc.add_paragraph()
                    add_runs_to_paragraph(subtitle, next_non_empty[1])
                    if subtitle.runs:
                        subtitle.runs[0].font.color.rgb = MUTED
                        subtitle.runs[0].font.size = Pt(11)
                    set_paragraph_spacing(subtitle, after=14)
                    i = next_non_empty[0]
            else:
                add_heading(doc, title_text, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], level=2)
            i += 1
            continue

        # 4. Bullet list items
        bullet_match = re.match(r'^[\-\*]\s+(.*)$', stripped)
        if bullet_match:
            item_text = bullet_match.group(1)
            # Check for checklists e.g., - [ ] or - [x]
            chk_match = re.match(r'^\[([ x\/])\]\s+(.*)$', item_text)
            if chk_match:
                status = chk_match.group(1)
                content = chk_match.group(2)
                box = "☐"
                if status == 'x':
                    box = "☑"
                elif status == '/':
                    box = "☒"
                item_text = f"{box} {content}"

            p = doc.add_paragraph(style="List Bullet")
            add_runs_to_paragraph(p, item_text)
            set_paragraph_spacing(p, after=4)
            i += 1
            continue

        # 5. Numbered / Step list items
        num_match = re.match(r'^\d+\.\s+(.*)$', stripped)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_runs_to_paragraph(p, num_match.group(1))
            set_paragraph_spacing(p, after=4)
            i += 1
            continue

        # 6. Standard paragraphs or blockquotes (which can span multiple lines)
        para_lines = []
        is_blockquote = stripped.startswith(">")
        
        while i < len(lines):
            l = lines[i].rstrip('\n')
            s = l.strip()
            if not s:
                break
            # Check if this line starts a new block element
            is_new_block = (
                s.startswith("#") or 
                re.match(r'^[\-\*]\s+', s) or 
                re.match(r'^\d+\.\s+', s) or 
                s in ["---", "***", "___"] or 
                s.startswith("![") or
                (s.startswith(">") != is_blockquote)
            )
            if is_new_block:
                break
            
            if is_blockquote:
                # Strip leading blockquote character
                s_clean = s.lstrip("> ").strip()
                para_lines.append(s_clean)
            else:
                para_lines.append(s)
            i += 1

        if para_lines:
            para_text = " ".join(para_lines)
            p = doc.add_paragraph()
            if is_blockquote:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                set_paragraph_spacing(p, before=8, after=8)
                add_runs_to_paragraph(p, para_text)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = MUTED
            else:
                p.style = doc.styles["Body Text"]
                add_runs_to_paragraph(p, para_text)
                set_paragraph_spacing(p)

    doc.save(OUT)
    print(f"Successfully compiled document to {OUT}")


if __name__ == "__main__":
    build()
